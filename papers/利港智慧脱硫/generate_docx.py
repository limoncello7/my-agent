#!/usr/bin/env python3
"""Generate 利港智慧脱硫 paper following the provided template format."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

doc = Document()

# ============================================================
# Page setup: A4
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_two_columns(section):
    """Set two-column layout on a section."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sect_pr.append(cols)
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '480')  # space between columns in twips (~0.85cm)


def add_column_break():
    """Insert a column break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(OxmlElement('w:br'))
    run._element[-1].set(qn('w:type'), 'column')


def set_font(run, name_cn='宋体', name_en='Times New Roman', size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def set_pf(p, space_before=0, space_after=0, line_spacing=1.25,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)


def add_title(text):
    """小二黑体(18pt), centered."""
    p = doc.add_paragraph()
    set_pf(p, space_before=12, space_after=6, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=18, bold=True)


def add_author(text):
    """四号楷体(14pt), centered."""
    p = doc.add_paragraph()
    set_pf(p, space_before=0, space_after=3, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_font(run, name_cn='楷体', size=14)


def add_affiliation(text):
    """五号宋体(10.5pt), centered."""
    p = doc.add_paragraph()
    set_pf(p, space_before=0, space_after=6, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_font(run, name_cn='宋体', size=10.5)


def add_abstract_label():
    """摘要 label + body: 小五号宋体(9pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=6, space_after=0, line_spacing=1.25,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=18)
    run_label = p.add_run('摘  要：')
    set_font(run_label, name_cn='宋体', size=9, bold=False)
    return p


def add_abstract_body(p, text):
    run_body = p.add_run(text)
    set_font(run_body, name_cn='宋体', size=9)


def add_keywords():
    """关键词: 小五号宋体(9pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=0, space_after=9, line_spacing=1.25,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=18)
    run = p.add_run('关键词：智慧脱硫；广义预测控制；iFGD；湿法脱硫；闭环控制；节能降耗')
    set_font(run, name_cn='宋体', size=9)


def add_heading0(text):
    """0 引言 style heading: 小四号黑体(12pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=9, space_after=3, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=12, bold=True)


def add_heading1(text):
    """一级标题 (1, 2...): 小四号黑体(12pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=9, space_after=3, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=12, bold=True)


def add_heading2(text):
    """二级标题 (1.1, 1.2...): 五号黑体(10.5pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=3, space_after=2, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=10.5, bold=True)


def add_heading3(text):
    """三级标题 (1.1.1...): 五号宋体(10.5pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=2, space_after=1, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_font(run, name_cn='宋体', size=10.5, bold=False)


def add_body(text):
    """正文: 五号宋体(10.5pt), first-line indent 2 chars (~21pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=0, space_after=0, line_spacing=1.25,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=21)
    run = p.add_run(text)
    set_font(run, name_cn='宋体', size=10.5)


def add_ref(text):
    """参考文献: 五号宋体(10.5pt)."""
    p = doc.add_paragraph()
    set_pf(p, space_before=0, space_after=0, line_spacing=1.25,
           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = p.add_run(text)
    set_font(run, name_cn='宋体', size=10.5)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, name_cn='黑体', size=9, bold=True)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_font(run, name_cn='宋体', size=9)


# ============================================================
# BUILD DOCUMENT
# ============================================================
add_title('利港智慧脱硫系统研究与分析')
add_author('董祥晖¹，任政达²')
add_affiliation('（浙江浙能科技环保集团股份有限公司，浙江省杭州市 310000）')

# Abstract
p_abs = add_abstract_label()
add_abstract_body(p_abs,
    '随着国家"双碳"战略目标的深入推进和火电机组超低排放标准的日趋严格，燃煤电厂脱硫系统面临'
    '着环保指标闭环控制不完善、设备动态特性耦合加剧、协同控制策略不明确、系统能耗偏高等多重'
    '挑战。本文以江阴利港2×100万千瓦机组扩建项目脱硫EPC总承包工程为研究对象，系统分析了其'
    '智慧脱硫控制系统（iFGD）的技术架构、核心算法与实际应用效果。iFGD系统采用外挂式架构，'
    '通过标准的MODBUS协议与机组DCS系统实现安全可靠的双向通讯。系统核心采用广义预测控制（GPC）'
    '算法，基于受控自回归积分滑动平均（CARIMA）模型，通过Diophantine方程递推求解和滚动优化'
    '策略，有效克服了湿法脱硫过程中吸收塔浆液pH值大滞后、大惯性的控制难题。实际运行数据表明：'
    '在稳态负荷工况下，出口SO₂浓度可稳定控制在设定值±5 mg/Nm³以内；在AGC变负荷工况下控制'
    '精度同样满足±5 mg/Nm³的要求；系统自动投运率达到100%。本文还深入分析了iFGD系统的优势'
    '与不足，并从算法优化、系统集成、数据智能和硬件可靠性四个维度提出了具体的改良方向。'
)
add_keywords()

# Insert continuous section break for two-column body
new_section = doc.add_section()
new_section.start_type = 0  # continuous section break
# Copy page setup from first section
for key in ['page_width', 'page_height', 'top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
    setattr(new_section, key, getattr(doc.sections[0], key))
set_two_columns(new_section)

# ============================================================
# 0 引言  (maps to original Chapter 1 content)
# ============================================================
add_heading0('0  引言')

add_heading2('0.1  研究背景')
add_body(
    '随着"碳达峰、碳中和"战略目标的持续推进，我国能源结构正在发生深刻变革。'
    '新能源占比不断提高，火电机组从传统的基荷运行逐步转向深度调峰运行模式，'
    '机组负荷变动频繁、变化幅度大，对环保设施的适应性和控制品质提出了更高要求。'
    '与此同时，国家环保排放标准日趋严格，超低排放改造已完成全面覆盖，'
    '火电企业面临着多重压力：既要确保环保指标稳定达标以避免经济处罚和社会形象损失，'
    '又要通过优化运行策略实现节能降耗以降低运营成本。'
)
add_body(
    '脱硫系统作为火电机组环保设施的核心环节，其运行水平直接关系到企业的环保达标和经营效益。'
    '然而，当前国内大量火电机组脱硫系统存在以下共性问题：第一，环保排放指标闭环控制逻辑'
    '不完善，出口SO₂浓度控制仍依赖运行人员手动调节，控制品质波动大；第二，脱硫系统内部'
    '浆液循环泵、供浆调节阀、氧化风机等多设备间的动态特性耦合加剧，导致协同控制策略不明确，'
    '难以实现全局最优；第三，脱硫过程连续调节能力不足，在负荷和煤质快速变化时容易导致排放'
    '超标；第四，由于控制手段落后，运行人员为保证环保达标不得不采用保守策略，导致脱硫系统'
    '能耗普遍偏高，存在巨大的节能降耗空间。'
)
add_body(
    '在上述背景下，智慧脱硫技术应运而生。江阴利港2×100万千瓦机组扩建项目所采用的iFGD'
    '（Intelligent Flue Gas Desulfurization）智慧脱硫控制系统，正是这一技术方向的典型代表。'
)

add_heading2('0.2  研究目的与意义')
add_body(
    '本文以江阴利港电厂#9、#10机组智慧脱硫控制系统项目为研究对象，系统梳理iFGD系统的技术架构、'
    '核心算法和实现方案，基于实际运行数据评估控制品质与综合效益，分析技术优势与不足，并提出改良方向。'
    '本研究旨在为火电行业脱硫系统的智能化改造提供可复制、可推广的技术参考，助力"双碳"目标实现。'
)

add_heading2('0.3  国内外研究现状')
add_body(
    '传统的石灰石-石膏湿法脱硫过程具有大滞后、大惯性和非线性时变特性，常规PID控制器难以在全'
    '工况范围内获得满意的控制品质。自Clarke等人于1987年提出广义预测控制（GPC）理论以来[4-5]，'
    '该算法在工业过程控制领域得到广泛应用，其多步预测、滚动优化和反馈校正的特征特别适合处理'
    '大时滞、慢动态过程。在国内，席裕庚[9]、诸静[10]等学者在预测控制领域做出了重要贡献。'
    '上海鉴智软件技术有限公司率先将GPC算法应用于火电机组脱硫控制，开发了iFGD智慧脱硫控制系统，'
    '在江阴利港电厂成功投运。然而，现有文献对智慧脱硫系统的系统性总结仍然较少，本文正针对'
    '这一缺口展开工作。'
)

add_heading2('0.4  论文结构安排')
add_body(
    '本文共分五章：本章为引言；第1章介绍项目概况与脱硫系统现状；第2章分析iFGD系统技术架构与'
    '核心算法；第3章基于运行数据评价实施效果；第4章总结技术优势、指出不足并提出改良方向；'
    '第5章给出结论与展望。'
)

# ============================================================
# 1 项目概况与脱硫系统现状分析
# ============================================================
add_heading1('1  项目概况与脱硫系统现状分析')

add_heading2('1.1  江阴利港电厂项目概况')
add_body(
    '江阴利港电厂2×100万千瓦机组扩建项目是大型火力发电建设工程，建设两台百万千瓦级超超临界'
    '燃煤发电机组（编号#9、#10）。根据国家环保要求，项目同步配套建设了高效烟气脱硫设施，'
    '采用石灰石-石膏湿法烟气脱硫（Wet Flue Gas Desulfurization, WFGD）工艺。'
)
add_body(
    '该脱硫系统的基本配置为：1炉1塔方案，每台机组脱硫系统配置5台浆液循环泵，设计工况下'
    '4用1备，校核工况下5台全部投用；配置5台高速离心式氧化风机，出口母管连通，设计工况下'
    '2用3备，校核工况下3用2备。每台机组脱硫氧化风管设有流量测试和调节阀，可对氧化风量进行'
    '独立调节控制。'
)
add_body(
    '智慧脱硫控制系统项目由上海鉴智软件技术有限公司承担，包括#9、#10机组"智慧脱硫控制系统"'
    '的研发、设计、安装、调试及验收等全部工作。项目于2026年初完成调试并交付运行投运，目前已'
    '实现烟囱净烟气出口SO₂浓度的闭环控制、浆液循环泵的节能指导以及氧化风量的自动优化控制。'
)

add_heading2('1.2  传统脱硫控制系统存在的问题')
add_body(
    '在iFGD系统投运前，利港电厂#9、#10机组脱硫系统存在以下问题：（1）控制方式落后，采用人工'
    '操作模式，控制品质依赖运行人员经验；（2）出口SO₂浓度未实现闭环控制，仍停留在开环手动调节'
    '层面；（3）过度保守运行导致能耗偏高，运行人员为保达标将SO₂浓度控制在远低于排放限值的水平，'
    '造成石灰石和电力的浪费；（4）超标风险不容忽视，负荷快速变化时手动控制响应滞后，易致排放'
    '超标；（5）运行人员需持续监控DCS画面并频繁操作，劳动强度大；（6）氧化风控制粗放，过量供风'
    '不仅增加电耗，还会导致浆液溢流等问题。'
)

add_heading2('1.3  智慧脱硫系统建设目标')
add_body(
    '针对上述问题，江阴利港电厂启动了智慧脱硫控制系统建设项目，设定了明确的热态验收目标：'
    '（1）实现脱硫系统吸收塔pH值、出口SO₂浓度的设定值可调闭环控制，以及氧化风量的合理控制，'
    '使运行人员能够像操作"自动驾驶"一样方便地管理脱硫系统；（2）在智慧脱硫优化系统运行期间，'
    '自动投入率不低于99.9%，确保系统能够长期连续稳定运行；（3）服务器负荷率不大于40%，保证'
    '控制系统本身不会对现有DCS系统造成额外负担；（4）在智慧脱硫优化系统运行期间，脱硫出口'
    'SO₂波动范围不超过±5 mg/Nm³，吸收塔pH值波动范围不超过±0.2。'
)

# ============================================================
# 2 iFGD智慧脱硫系统技术架构与核心算法
# ============================================================
add_heading1('2  iFGD智慧脱硫系统技术架构与核心算法')

add_heading2('2.1  系统总体架构')
add_body(
    'iFGD系统采用外挂式架构设计，这一设计理念的核心原则是：在任何情况下都不影响现有DCS系统的'
    '安全性和稳定性。系统通过标准的MODBUS协议与电厂DCS系统实现双向通讯，既可以从DCS采集所需'
    '的过程变量实时数据，也可以将优化控制指令安全地传输给DCS执行。'
)
add_body(
    'iFGD系统由三大核心模块组成：（1）数据采集与通讯模块。该模块负责通过MODBUS/OPC协议从'
    'DCS系统实时采集脱硫过程的关键参数，包括入口/出口SO₂浓度、烟气量、浆液循环泵电流、供浆'
    '流量、吸收塔pH值、机组负荷等。同时，该模块将控制器计算出的最优操作值进行严格的数据校验'
    '和平滑处理后，通过MODBUS协议写入DCS系统执行。（2）智能控制算法模块。这是iFGD系统的核心，'
    '集成了广义预测控制（GPC）算法、多变量动态关联模型、浆液循环泵组合优化算法和氧化风量智能'
    '寻优算法。该模块通过对脱硫FGD过程进行实时动态建模，建立多变量动态关联模型，并基于该模型'
    '进行每时每刻的优化控制，保证脱硫控制品质。（3）可视化监控模块。系统提供了丰富的可视化和'
    '组态功能，包括运行参数实时显示、实时/历史趋势曲线、控制性能在线评估、经济指标实时统计、'
    '时均值统计分析等界面。系统还基于influxDB数据库实现了历史数据存储统计与能耗自动建模功能。'
)

add_heading2('2.2  广义预测控制（GPC）算法')
add_body(
    '广义预测控制（GPC）是iFGD系统的核心算法，用于解决脱硫吸收塔pH值这一关键参数的控制问题。'
    '湿法脱硫过程中，吸收塔浆液pH值对脱硫效率有决定性影响，但pH值对供浆流量的响应具有明显的'
    '大滞后、大惯性和慢动态特性，传统的PID控制器很难获得理想的控制效果。'
)
add_body(
    'GPC算法采用受控自回归积分滑动平均（CARIMA）模型来描述脱硫过程的动态特性。该模型可以'
    '表示为：A(q⁻¹)yₜ = B(q⁻¹)uₜ₋d + ξₜ/Δ，其中yₜ为系统输出（如pH值或SO₂浓度），'
    'uₜ为控制量（如供浆阀门开度），ξₜ为随机噪声，Δ为差分算子，d为系统纯时延步数。'
)
add_body(
    'GPC算法的核心机制包括：多步预测（基于CARIMA模型和Diophantine方程递推预测未来N个时刻的'
    '系统输出）、滚动优化（定义包含预测偏差和控制增量的目标函数，采用柔化跟踪轨迹技术避免控制量'
    '剧烈波动）、反馈校正（利用实时反馈补偿模型不确定性和外部扰动）。GPC算法计算快速、鲁棒性强，'
    '此外系统还设计了基于pH值、烟气温度等因素的模型在线修正功能，进一步提升了控制品质。'
)

add_heading2('2.3  多变量协同优化控制策略')
add_body(
    '湿法脱硫系统是一个典型的多变量耦合系统。iFGD系统基于物料平衡机理分析，建立SO₂排放与运行'
    '参数的动态控制模型，实现了以下协同优化功能：（1）供浆流量与pH值的协同控制，通过GPC算法'
    '实时优化供浆流量，减少石灰石消耗；（2）浆液循环泵组合优化，以能耗最优为目标计算最优泵组合'
    '方案；（3）氧化风量智能调节，建立以进口烟气量和脱硫效率为输入的优化模型，降低电耗并避免'
    '浆液溢流；（4）多目标动态寻优，建立机理与数据融合的能耗模型，实现运行方式的多目标优化。'
)

add_heading2('2.4  时序多层控制架构')
add_body(
    '针对超低排放考核中"小时均值"这一关键指标，iFGD系统创新性地采用了时序多层控制架构：上层'
    '实时计算累积小时均值，动态评估超标风险并在必要时调整下层设定值；下层基于GPC算法追踪上层'
    '设定值，通过精确调节供浆流量实现实时值的精细控制。该架构突破了仅关注瞬时值的局限，将小时'
    '均值直接纳入控制目标体系，从源头上避免了超标风险。'
)

add_heading2('2.5  三种运行控制模式')
add_body(
    '为适应不同的运行工况和需求场景，iFGD系统设计了三种控制模式，运行人员可根据实际情况灵活'
    '切换。SO₂模式（推荐默认模式）以净烟气SO₂浓度为控制目标，运行人员只需设置SO₂设定值，系统'
    '在设定的pH和供浆流量上下限范围内自动调整pH设定值和供浆量，适用于日常正常运行。pH模式以'
    '吸收塔pH值为控制目标，运行人员设置pH设定值，系统在供浆流量上下限范围内自动调整供浆量，'
    '适用于出口SO₂测量异常或调试/试验工况。流量模式以供浆流量为控制目标，运行人员设置流量设定值，'
    '系统自动调整供浆阀门开度，适用于调试或试验工况。三种控制模式的具体对比如表1所示。'
)
add_table(
    ['控制模式', '控制目标', '工作原理', '适用场景'],
    [
        ['SO₂模式\n（推荐默认）', '净烟气SO₂浓度',
         '设置SO₂设定值，系统自动调整pH设定值和供浆量', '日常正常运行'],
        ['pH模式', '吸收塔pH值',
         '设置pH设定值，系统自动调整供浆量', '出口SO₂测量异常或调试工况'],
        ['流量模式', '供浆流量',
         '设置流量设定值，系统自动调整供浆阀门开度', '调试或试验工况'],
    ]
)
# add table caption
p = doc.add_paragraph()
set_pf(p, space_before=0, space_after=3, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
run = p.add_run('表1  三种运行控制模式对比')
set_font(run, name_cn='宋体', size=9, bold=False)

add_heading2('2.6  安全策略与系统可靠性设计')
add_body(
    '系统设计了多层次安全策略：心跳信号机制（异常时自动切除iFGD输出、无缝切换回原控制模式）、'
    '数据验证与保护（对输入输出数据进行严格校验和平滑处理）、限值保护（设定了供浆流量、pH值等'
    '安全限值）以及二次开发与组态功能（基于JSP语法的配置文件系统，支持非程序人员修改规则）。'
)

# ============================================================
# 3 项目实施效果分析
# ============================================================
add_heading1('3  项目实施效果分析')

add_heading2('3.1  稳态与变负荷工况控制效果')
add_body(
    '在稳定负荷工况下，iFGD系统展现了优异的控制品质。以#9机组2026年3月6日运行为例：出口SO₂'
    '浓度设定值8 mg/Nm³，实际平均值8.2 mg/Nm³，标准方差仅3 mg/Nm³，波动幅度极小。'
)
add_body(
    '在AGC变负荷工况下，系统表现同样突出。2026年3月18日降负荷过程（1000MW→400MW）中，'
    '出口SO₂浓度稳定控制在设定值±5 mg/Nm³以内，设定值5 mg/Nm³，实际平均值4.3 mg/Nm³。'
    '2026年3月10日升负荷过程（310MW→810MW）中，设定值10 mg/Nm³，实际平均值11.2 mg/Nm³，'
    '同样满足±5 mg/Nm³的控制精度。设定值阶跃响应平稳，过渡时间约35-40分钟，超调量小于'
    '1 mg/Nm³。GPC算法的预测能力使其能够在扰动影响输出前进行补偿调节，有效克服了脱硫过程'
    '的大滞后特性。'
)

add_heading2('3.2  系统投运率与可靠性')
add_body(
    '在2026年3月连续运行期间，#9、#10机组iFGD系统除去因pH计标定、CEMS标定等客观原因需'
    '暂时切回手动外，自动投运率达100%，远超项目设定的99.9%目标。服务器CPU负荷率低于40%，'
    '系统具有良好的计算效率。可视化监控界面实时显示控制回路运行信息，经济指标统计界面可计算'
    '石灰石消耗量、脱硫电耗和单位脱硫成本等关键参数，为运行管理提供了量化决策支持。'
)

add_heading2('3.3  综合效益分析')
add_body(
    'iFGD系统的经济效益体现在：石灰石消耗量降低（GPC精确控制供浆流量避免过量投加）；脱硫系统'
    '电耗减少（循环泵组合优化和氧化风量智能调节）；环保超标风险消除（闭环控制避免排放超标，'
    '确保电价补贴全额获取）；运行维护成本降低。环保方面，系统确保了超低排放稳定达标，同时间接'
    '减少了石灰石开采和运输过程的碳排放。社会效益方面，运行人员劳动强度大幅降低，实现了'
    '"环保、节能、自动"的有机融合。'
)

# ============================================================
# 4 智慧脱硫技术优势与改良方向探讨
# ============================================================
add_heading1('4  智慧脱硫技术优势与改良方向探讨')

add_heading2('4.1  iFGD系统的核心优势分析')

add_heading3('4.1.1  技术层面优势')
add_body(
    'iFGD系统在技术层面的核心优势体现在：第一，先进预测控制算法。GPC通过多步预测和滚动优化'
    '实现了"事前控制"，有效克服了大滞后难题，与PID的"事后调节"有本质区别。第二，外挂式架构设计。'
    '系统不修改DCS内部逻辑，通过MODBUS通讯实现数据交互，最大程度降低了对现有系统的安全风险。'
    '第三，多变量协同优化。将pH值、供浆流量、循泵组合等多个参数纳入统一优化框架，实现了从'
    '"局部最优"到"全局最优"的跨越。第四，时序多层控制。通过实时值与小时均值双层架构，兼顾了'
    '短期排放稳定和长期考核达标。'
)

add_heading3('4.1.2  运行与经济层面优势')
add_body(
    '在运行层面，iFGD实现了从人工操作到一键智能控制、从经验驱动到数据驱动、从被动响应到主动预测'
    '的"三个转变"。运行人员只需点击"脱硫优化投入"按钮即可完成全部控制任务。经济方面，通过循环泵'
    '组合优化、供浆量精确控制和氧化风量合理调节，系统在降低石灰石消耗和电耗方面成效显著，且接近'
    '100%的投运率确保了这些效益的持续性。'
)

add_heading2('4.2  iFGD系统与传统PID控制的对比')
add_body(
    '为客观评价iFGD系统的技术先进性，将其与传统PID控制方案进行系统对比，结果如表2所示。'
)
add_table(
    ['对比维度', '传统PID控制', 'iFGD预测控制'],
    [
        ['大滞后处理能力', '困难，易产生超调和振荡', '优秀，通过预测补偿实现平稳控制'],
        ['多变量耦合处理', '无法有效解耦，各回路相互干扰', '多变量协同优化，全局最优'],
        ['鲁棒性和适应性', '参数固定，工况变化时需重新整定', '模型在线修正，自适应能力强'],
        ['控制精度', '±10~15 mg/Nm³', '±5 mg/Nm³以内'],
        ['自动化程度', '需人工频繁干预调整', '全自动闭环运行，一键投入/切除'],
        ['对运行人员的要求', '需要丰富经验和持续关注', '操作简便，大幅降低技能门槛'],
        ['节能降耗能力', '有限，保守运行导致能耗偏高', '显著，多目标优化实现成本最优'],
        ['系统可扩展性', '受限于DCS硬件和逻辑', '外挂式架构，灵活扩展升级'],
    ]
)
p = doc.add_paragraph()
set_pf(p, space_before=0, space_after=3, line_spacing=1.25, alignment=WD_ALIGN_PARAGRAPH.CENTER)
run = p.add_run('表2  iFGD系统与传统PID控制对比')
set_font(run, name_cn='宋体', size=9, bold=False)

add_heading2('4.3  当前技术存在的不足')
add_body(
    '尽管iFGD系统取得了显著成功，但仍存在以下不足：（1）对CEMS和pH计等在线仪表的依赖度高，'
    '仪表故障直接影响系统可用率；（2）GPC算法在极端煤质变化或深度调峰工况下模型精度可能下降；'
    '（3）GPC参数整定需较专业的控制理论知识，增加了调试运维门槛；（4）系统主要专注于脱硫控制，'
    '与其他环保设施（如脱硝、除尘）的协同控制尚未完全打通。'
)

add_heading2('4.4  改良方向与优化建议')
add_body(
    '面向未来，智慧脱硫技术可从以下维度持续升级：（1）算法层面，引入LSTM或Transformer等深度学习'
    '模型与GPC框架结合，研发基于在线学习的自适应算法；（2）系统集成层面，将脱硫、脱硝、除尘纳入'
    '统一智能环保控制平台，构建数字孪生系统；（3）数据与智能层面，引入强化学习实现控制策略自优化，'
    '开发AI智能故障诊断与预测性维护功能；（4）硬件可靠性层面，对关键仪表采用冗余配置，开发在线'
    '自动标定技术，加强工控网络安全防护。'
)

# ============================================================
# 5 结论与展望
# ============================================================
add_heading1('5  结论与展望')

add_heading2('5.1  主要结论')
add_body(
    '本文以江阴利港2×100万千瓦机组扩建项目脱硫EPC总承包工程的iFGD智慧脱硫控制系统为研究对象，'
    '系统分析了其技术架构、核心算法和实施效果。通过本研究可以得出以下主要结论：'
)
add_body(
    '（1）iFGD系统基于GPC算法，采用外挂式架构和MODBUS通讯协议，有效解决了湿法脱硫过程中pH值'
    '大滞后、大惯性的控制难题，实现了全自动闭环优化控制。（2）实际运行数据充分验证了系统品质：'
    '稳态工况下出口SO₂浓度偏差小于±5 mg/Nm³，标准方差仅3 mg/Nm³；AGC变负荷工况下仍保持±5 mg/Nm³'
    '的控制精度；设定值阶跃响应平稳，过渡时间约30-40分钟。（3）系统自动投运率达100%（排除仪表标定'
    '等客观因素），超过99.9%的项目目标。（4）系统通过浆液循环泵组合优化、氧化风量智能调节实现了'
    '显著的节能降耗效果，大幅降低了运行人员操作强度。（5）与传统PID控制相比，iFGD系统在大滞后处理、'
    '多变量耦合处理、自适应能力和控制精度等指标上具有显著优势。'
)

add_heading2('5.2  技术创新点总结')
add_body(
    '本项目技术创新主要体现在四个方面：一是GPC算法在脱硫pH值控制中的工程化应用，实现了从理论到'
    '工程的完整转化；二是时序多层控制架构设计，将小时均值纳入控制目标体系；三是机理与数据驱动融合的'
    '建模方法，构建了兼顾物理约束和数据特征的混合模型；四是外挂式架构在存量机组改造中的适配性创新，'
    '为存量机组智能化改造提供了可行的技术路径。'
)

add_heading2('5.3  推广前景与应用展望')
add_body(
    'iFGD系统的成功投运验证了系统的可靠性、安全性和经济性。其外挂式架构和标准化通讯协议可灵活适配'
    '不同容量等级（30万~100万千瓦）和不同控制系统的机组，为已完成超低排放改造的存量机组提供了'
    '见效快、风险低、效益高的技术升级路径。展望未来，智慧脱硫技术将沿着"智能化→协同化→自主化"的'
    '路径演进，逐步实现"无人干预、自适应优化"的终极目标。'
)

# ============================================================
# 参考文献
# ============================================================
add_heading1('参考文献')
refs = [
    '[1] 上海鉴智软件技术有限公司. 江阴利港2×100万千瓦机组扩建项目脱硫EPC总承包工程智慧脱硫控制系统总结报告[R]. 2026.',
    '[2] 上海鉴智软件技术有限公司. 超低排放FGD装置净烟气SO₂浓度智能闭环控制系统iFGD工程师维护手册[Z]. 2026.',
    '[3] 上海鉴智软件技术有限公司. 脱硫系统优化控制系统(iFGD)运行操作手册[Z]. 2026.',
    '[4] Clarke D W, Mohtadi C, Tuffs P S. Generalized predictive control—Part I. The basic algorithm[J]. Automatica, 1987, 23(2): 137-148.',
    '[5] Clarke D W, Mohtadi C, Tuffs P S. Generalized predictive control—Part II. Extensions and interpretations[J]. Automatica, 1987, 23(2): 149-160.',
    '[6] Richalet J, Rault A, Testud J L, et al. Model predictive heuristic control: Applications to industrial processes[J]. Automatica, 1978, 14(5): 413-428.',
    '[7] Qin S J, Badgwell T A. A survey of industrial model predictive control technology[J]. Control Engineering Practice, 2003, 11(7): 733-764.',
    '[8] Camacho E F, Bordons C. Model Predictive Control[M]. 2nd ed. London: Springer-Verlag, 2004.',
    '[9] 席裕庚. 预测控制[M]. 北京: 国防工业出版社, 2013.',
    '[10] 诸静. 智能预测控制及其应用[M]. 杭州: 浙江大学出版社, 2002.',
    '[11] 国家发展和改革委员会, 国家能源局. "十四五"现代能源体系规划[Z]. 2022.',
    '[12] 生态环境部. 关于推进实施水泥行业超低排放的意见[Z]. 2024.',
    '[13] 中国电力企业联合会. 中国电力行业年度发展报告2025[R]. 2025.',
    '[14] 刘永文, 刘吉臻, 等. 火电机组脱硫系统建模与优化控制研究综述[J]. 中国电机工程学报, 2023, 43(15): 5867-5885.',
    '[15] 张兴, 李海涛, 等. 基于预测控制的湿法脱硫pH值控制策略研究[J]. 电力科技与环保, 2022, 38(4): 308-315.',
]
for ref in refs:
    add_ref(ref)

# ============================================================
# Save
# ============================================================
output_path = r'C:\Users\D1405\my-agent\papers\利港智慧脱硫\利港智慧脱硫系统研究与分析.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
